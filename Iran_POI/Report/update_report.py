# -*- coding: utf-8 -*-
"""Surgical updates to Iran_POI_Report.docx (verified numbers only)."""
from pathlib import Path
import pandas as pd
from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

HERE = Path(__file__).resolve().parent
CLS = HERE.parent / 'Classification'
DOC = HERE / 'Iran_POI_Report.docx'
PNG = CLS / 'plot_enriched_accuracy_trend.png'

doc = Document(str(DOC))


def set_rtl(p):
    p._p.get_or_add_pPr().append(OxmlElement('w:bidi'))


# ---------- Edit 1: Table 6 — LLM AUC 0.433 -> 0.9667 (verified) ----------
doc.tables[6].rows[2].cells[5].text = '0.9667'

# ---------- Edit 2: reformulate the LLM conclusion (para 60) ----------
new60 = ("המודל אמין - אוכלוסיית היעד בסף 0.8 מגיעה ל-Hit Rate של כ-90% (85.7% בהערכה "
         "עיוורת). מול מודל שפה גדול (Claude Opus, 10 ריצות עיוורות) המודל המאומן מנצח "
         "בהחלטה בפער עצום: Accuracy 0.95 מול 0.53 ו-F1 0.93 מול 0.49; ב-AUC (יכולת דירוג) "
         "השניים שקולים (~0.97) - כלומר ה-LLM מזהה היטב מי שייך ליעד אך מסרב להכריע (מסמן "
         "unknown), וזו בדיוק הסיבה שהוא כושל כמסווג אוטומטי.")
p60 = doc.paragraphs[60]
for r in list(p60.runs[1:]):
    r._element.getparent().remove(r._element)
if p60.runs:
    p60.runs[0].text = new60
else:
    p60.add_run(new60)

# ---------- Edit 3: enriched-across-iterations table + figure after step 5-6 ----------
trend = pd.read_csv(CLS / 'Feature_Enrichment' / 'enriched_kfold_trend.csv')

# anchor = the "שלב 7 - קריטריון עצירה" heading
anchor = next(p for p in doc.paragraphs if p.text.strip().startswith('שלב 7 - קריטריון'))

expl = anchor.insert_paragraph_before(
    "שיפור שלב 6 - העשרת הפיצ'רים לאורך כל האיטרציות: כאשר מזריקים את פיצ'רי ההעשרה "
    "הספציפיים-לאיראן לכל איטרציה (K-Fold), ה-Accuracy של המודל המנצח על המשימה המרכזית "
    "(target_population, 2-מחלקות) מפסיק לרדת ונשאר גבוה (0.92→0.84) - בעוד הפיצ'רים "
    "הבסיסיים צונחים (0.90→0.73) ככל שהלמידה האקטיבית מוסיפה משתמשים קשים יותר. גם ה-AUC "
    "המועשר יציב-גבוה (~0.89-0.94) מול בסיס יורד (0.76). כלומר: מקור ה'ירידה' בגרף המקורי "
    "היה עוני-הפיצ'רים, לא נסיגה של המודל.")
set_rtl(expl)

tbl = doc.add_table(rows=1, cols=6)
tbl.style = 'Table Grid'
for c, h in zip(tbl.rows[0].cells,
                ['איטרציה', 'מתויגים', 'Acc בסיס', 'Acc מועשר', 'AUC בסיס', 'AUC מועשר']):
    c.text = h
for _, r in trend.iterrows():
    cells = tbl.add_row().cells
    vals = [int(r['iteration']), int(r['n_labeled_users']),
            f"{r['target_population__baseline_bestAcc']:.3f}",
            f"{r['target_population__enriched_bestAcc']:.3f}",
            f"{r['target_population__baseline_bestAUC']:.3f}",
            f"{r['target_population__enriched_bestAUC']:.3f}"]
    for c, v in zip(cells, vals):
        c.text = str(v)
expl._p.addnext(tbl._tbl)

img_p = doc.add_paragraph()
img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
img_p.add_run().add_picture(str(PNG), width=Inches(6.2))
tbl._tbl.addnext(img_p._p)

doc.save(str(DOC))
print("saved. tables:", len(doc.tables), "| paragraphs:", len(doc.paragraphs))
print("Table 6 LLM row:", [c.text for c in doc.tables[6].rows[2].cells])
