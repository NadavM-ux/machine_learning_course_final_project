# -*- coding: utf-8 -*-
"""Two surgical polish fixes. Run ONCE (not idempotent)."""
from pathlib import Path
from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Pt

HERE = Path(__file__).resolve().parent
DOC = HERE / 'Iran_POI_Report.docx'
doc = Document(str(DOC))


def set_rtl(p):
    p._p.get_or_add_pPr().append(OxmlElement('w:bidi'))


def set_text(p, text):
    for r in list(p.runs[1:]):
        r._element.getparent().remove(r._element)
    if p.runs:
        p.runs[0].text = text
    else:
        p.add_run(text)


# ---- Fix 1: reconcile 946 (rows) vs 945 (unique) ----
p946 = next(p for p in doc.paragraphs if '946 משתמשים מועמדים' in p.text)
set_text(p946,
         "946 פרופילים מועמדים (945 משתמשים ייחודיים; שדות הפרופיל: שם, ביו, מיקום, "
         "עוקבים/נעקבים/פוסטים, תאריך הצטרפות)")

# ---- Fix 2: explicit 0.8 threshold justification under 8-A ----
h8a = next(p for p in doc.paragraphs if p.text.strip().startswith('שלב 8-A'))
just = OxmlElement('w:p')
h8a._p.addnext(just)
from docx.text.paragraph import Paragraph
jp = Paragraph(just, h8a._parent)
set_rtl(jp)
run = jp.add_run(
    "בחירת הסף 0.8: בסף 0.9 אף משתמש אינו עובר (0 משתמשים) — מחמיר מדי; בסף 0.7 נכנסים 238 "
    "משתמשים אך הדיוק מדולל במקרים גבוליים; סף 0.8 מזהה 173 משתמשים תוך שמירה על Hit Rate של "
    "כ-90%, ולכן נבחר כאיזון הטוב ביותר בין כיסוי לדיוק.")
run.font.size = Pt(11)

doc.save(str(DOC))
print("Fix 1 (946/945):", p946.text[:60])
print("Fix 2 (0.8 justification) inserted after:", h8a.text.strip())
