# -*- coding: utf-8 -*-
"""Additive-only: gaps 1,3,4,5,6 (NOT 2). Adds new graph/tables/sentences at
section boundaries. Touches no existing paragraph/table/wording. Run once."""
from pathlib import Path
import pandas as pd
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

HERE = Path(__file__).resolve().parent
CLS = HERE.parent / 'Classification'
STEP7 = CLS / 'Step7_Stopping_Criteria'
STEP3 = CLS / 'Iteration_1' / 'Step3_Manual_Labeling'
DOC = HERE / 'Iran_POI_Report.docx'

# ---------- (1) generate a clean mean-Accuracy-by-iteration graph ----------
perf = pd.read_csv(STEP7 / 'stopping_criteria_performance_summary.csv')
ACC_PNG = STEP7 / 'plot_accuracy_by_iteration.png'
fig, ax = plt.subplots(figsize=(8, 4.2))
ax.plot(perf['iteration'], perf['mean_accuracy'], 'o-', color='#1f6f8b', lw=2)
for _, r in perf.iterrows():
    ax.annotate(f"{r['mean_accuracy']:.2f}", (r['iteration'], r['mean_accuracy']),
                textcoords="offset points", xytext=(0, 8), ha='center', fontsize=8)
ax.set_xlabel('Active-Learning iteration'); ax.set_ylabel('mean Accuracy (all classifiers)')
ax.set_title('Mean Accuracy per iteration (averaged over all classifiers)')
ax.set_ylim(0.4, 0.8); ax.grid(alpha=0.3)
ax.set_xticks(perf['iteration'])
fig.tight_layout(); fig.savefig(ACC_PNG, dpi=130); plt.close(fig)
print("wrote", ACC_PNG.name)

doc = Document(str(DOC))


def rtl(p):
    p._p.get_or_add_pPr().append(OxmlElement('w:bidi')); return p


def para(text, size=11, bold=False):
    p = doc.add_paragraph(); rtl(p)
    r = p.add_run(text); r.font.size = Pt(size); r.bold = bold
    return p


def cap(text):
    p = doc.add_paragraph(); rtl(p); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.italic = True; r.font.size = Pt(9)
    return p


def image_para(png, width=6.2):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(png), width=Inches(width))
    return p


def table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = 'Table Grid'
    for c, h in zip(t.rows[0].cells, headers):
        c.text = h
    for row in rows:
        cells = t.add_row().cells
        for c, v in zip(cells, row):
            c.text = str(v)
    return t


def find(prefix):
    return next(p for p in doc.paragraphs if p.text.strip().startswith(prefix))


def chain(anchor_para, *elements):
    """insert elements (paragraph/table objects created at end) right before anchor_para, in order."""
    cursor = anchor_para.insert_paragraph_before('', style=None)._p  # placeholder anchor spot
    for el in elements:
        node = el._p if hasattr(el, '_p') else el._tbl
        cursor.addnext(node); cursor = node
    # remove the empty placeholder
    cursor0 = cursor
    # placeholder is the first inserted; find and drop it
    ph = anchor_para._p.getprevious()
    # simpler: the placeholder was the very first 'cursor'; remove by locating empty para we made
    return


# ===== BLOCK A — gap 3: class-distribution tables (step 3), before "שלבים 5-6" =====
anchorA = find("שלבים 5-6")
capA = anchorA.insert_paragraph_before("התפלגות המחלקות בתיוג הידני (100 משתמשים, שלב 3):", style=None)
rtl(capA)
cursor = capA._p
CLASS_TASKS = [
    ("target_population", "iteration_1_target_population_summary.csv"),
    ("locals_vs_diaspora", "iteration_1_locals_vs_diaspora_summary.csv"),
    ("person_vs_organization", "iteration_1_person_vs_organization_summary.csv"),
]
for name, fn in CLASS_TASKS:
    df = pd.read_csv(STEP3 / fn)
    lbl = para(name, size=10, bold=True)
    cursor.addnext(lbl._p); cursor = lbl._p
    t = table(["class", "count", "percentage"],
              [[r['class'], int(r['count']), r['percentage']] for _, r in df.iterrows()])
    cursor.addnext(t._tbl); cursor = t._tbl

# ===== BLOCK B — gaps 1 & 4: summary table + accuracy graph + step-7 histogram, before "שלב 7+" =====
anchorB = find("שלב 7+")
capB = anchorB.insert_paragraph_before("טבלת סיכום ביצועים לפי איטרציה (ממוצע על פני כל המסווגים):", style=None)
rtl(capB)
cursor = capB._p
tB = table(["איטרציה", "מתויגים", "Accuracy ממוצע", "AUC ממוצע"],
           [[int(r['iteration']), int(r['n_labeled_users']),
             f"{r['mean_accuracy']:.3f}", f"{r['mean_AUC']:.3f}"] for _, r in perf.iterrows()])
cursor.addnext(tB._tbl); cursor = tB._tbl
g1 = image_para(ACC_PNG); cursor.addnext(g1._p); cursor = g1._p
c1 = cap("גרף Accuracy ממוצע לפי איטרציה, על פני כל המסווגים (שלב 6)."); cursor.addnext(c1._p); cursor = c1._p
h1 = image_para(STEP7 / 'stopping_criteria_confidence_histogram.png'); cursor.addnext(h1._p); cursor = h1._p
c2 = cap("היסטוגרמת רמות הביטחון על המשתמשים הלא-מתויגים (שלב 7)."); cursor.addnext(c2._p); cursor = c2._p

# ===== BLOCK C — gaps 5a, 5b, 6: three defense sentences, before "שלב 9" =====
anchorC = find("שלב 9")
s5a = anchorC.insert_paragraph_before(
    "הערה על ה-AUC של ה-LLM: מכיוון שהמודל מחזיר תווית קשיחה בכל הרצה, ה-AUC חושב משיעור "
    "ההצבעה למחלקת target על פני 10 ההרצות (fraction of runs), המשמש כ-P(target) לכל משתמש. "
    "כך הדירוג מוגדר, וזה בדיוק מה שחושף שה-LLM מדרג היטב אך מסרב להכריע.", style=None)
rtl(s5a)
s5b = anchorC.insert_paragraph_before(
    "מודל השפה שבו נעשה שימוש הוא Claude Opus (מודל חזית), המותר לפי הנחיית ה-PDF 'ChatGPT-5.2 "
    "או מודל גבוה יותר'. ההרצות בוצעו ב-16 ביולי 2026, 10 הרצות עיוורות והכרעת רוב.", style=None)
rtl(s5b)
s6 = anchorC.insert_paragraph_before(
    "היקף ההשוואה: n=19 משתמשי hold-out ודאיים בלבד, ולכן רווח הסמך ל-AUC רחב (בקירוב 0.75-0.99). "
    "ההשוואה המשמעותית היא הפער בין Accuracy ל-AUC, ולא הערך המוחלט.", style=None)
rtl(s6)

doc.save(str(DOC))
print("done. paragraphs:", len(doc.paragraphs), "| tables:", len(doc.tables),
      "| images:", sum(1 for r in doc.part.rels.values() if 'image' in r.reltype))
