# -*- coding: utf-8 -*-
"""Rewrite step-4 flowchart captions as flowing Hebrew (no arrows -> fixes RTL),
drop AI-looking em-dashes. Run once."""
from pathlib import Path
from docx import Document

DOC = Path(__file__).resolve().parent / 'Iran_POI_Report.docx'
d = Document(str(DOC))

INTRO = ("כדי לתעד באופן פורמלי את תהליך קבלת ההחלטות בתיוג הידני, בנינו שלושה תרשימי זרימה, "
         "אחד לכל משימת סיווג. כל תרשים מתאר את רצף השאלות שלפיו הוכרע כל משתמש, כאשר כל עלה "
         "מוביל לאחת התוויות. אותם תרשימים שימשו גם כהנחיה למודל השפה בשלב 8-B, כך שהמסווג "
         "המאומן ומודל השפה הוערכו לפי אותה לוגיקה. הקובץ המלא: user_labeling_decision_flows.pptx.")

CAP1 = ("תרשים 1, המשימה target_population: כך נקבע אם המשתמש שייך לאוכלוסיית היעד. אם הפרופיל "
        "מושהה, ריק או מוגן וללא מידע, הוא מסומן unknown. אם הביו, שם התצוגה או שם המשתמש מזכירים "
        "איראן, פרסית או עיר איראנית, הוא מסומן Target. אם שדה המיקום מצביע על עיר איראנית, גם אז "
        "Target. אם ציוצים אישיים מעידים שהאדם איראני, שוב Target. אחרת, כאשר קיימת זהות ברורה "
        "שאינה איראנית הוא מסומן Not Target, ובכל מקרה אחר נותר Unknown.")

CAP2 = ("תרשים 2, המשימה locals_vs_diaspora: ההבחנה בין מקומי לתפוצות, רלוונטית רק כאשר המשתמש "
        "שייך ליעד. אם שדה המיקום מציג עיר איראנית, הוא מסומן Local. אם הוא מציג עיר מחוץ לאיראן, "
        "הוא מסומן diaspora. אחרת, אם הביו או שם התצוגה מרמזים על תפוצה, הוא diaspora. אם עדיין "
        "לא הוכרע, ציוצים אחרונים שממקמים אותו בתוך איראן קובעים local, וכאלה שממקמים אותו מחוץ "
        "לאיראן קובעים diaspora. בהיעדר סימן ברור הוא נותר Unknown.")

CAP3 = ("תרשים 3, המשימה person_vs_organization: ההבחנה בין חשבון של אדם לחשבון של ארגון. אם "
        "הפרופיל מושהה, ריק או פרטי, הסיווג הוא Unknown. אם שם התצוגה או שם המשתמש מעידים בבירור "
        "על ארגון, הוא מסומן organization. אחרת, אם הביו מתאר מוסד, שוב organization. אם עדיין לא, "
        "כאשר תמונת הפרופיל היא לוגו או סמל והציוצים כתובים בקול מוסדי או בלשון רבים, הוא "
        "organization. לבסוף, אם יש ראיה ברורה לאדם יחיד הוא מסומן Person, ואחרת נותר Unknown.")


def set_para_text(p, text):
    for r in list(p.runs[1:]):
        r._element.getparent().remove(r._element)
    if p.runs:
        p.runs[0].text = text
    else:
        p.add_run(text)


def replace_in_para(p, old, new):
    for r in p.runs:
        if old in r.text:
            r.text = r.text.replace(old, new); return True
    full = ''.join(r.text for r in p.runs)
    if old in full and p.runs:
        for r in list(p.runs[1:]):
            r._element.getparent().remove(r._element)
        p.runs[0].text = full.replace(old, new); return True
    return False


def find(sub):
    return next(p for p in d.paragraphs if sub in p.text)


# whole-paragraph rewrites (remove arrows + numbered lists)
set_para_text(find("כדי לתעד באופן פורמלי"), INTRO)
set_para_text(find("שיוך לאוכלוסיית היעד"), CAP1)
set_para_text(find("מקומי מול תפוצות"), CAP2)
set_para_text(find("אדם מול ארגון"), CAP3)

# drop AI-looking em-dash in the 0.8 justification
replace_in_para(find("בחירת הסף 0.8"), " — מחמיר מדי", " והוא מחמיר מדי")

d.save(str(DOC))

# verify
d2 = Document(str(DOC))
arrows = [i for i,p in enumerate(d2.paragraphs) if '→' in p.text]
emdash = [i for i,p in enumerate(d2.paragraphs) if '—' in p.text]
print("paragraphs still with arrow →:", arrows)
print("paragraphs still with em-dash —:", emdash)
print("\n--- new captions ---")
for key in ["target_population:", "locals_vs_diaspora:", "person_vs_organization:"]:
    print("•", next(p.text for p in d2.paragraphs if key in p.text)[:90], "...")
