# -*- coding: utf-8 -*-
"""Style cleanup: drop bot-like (English) gloss parentheses + all separator lines. Run once."""
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn

DOC = Path(__file__).resolve().parent / 'Iran_POI_Report.docx'
d = Document(str(DOC))


def drop_borders(p):
    pPr = p._p.find(qn('w:pPr'))
    if pPr is not None:
        bdr = pPr.find(qn('w:pBdr'))
        if bdr is not None:
            pPr.remove(bdr)
            return True
    return False


def remove_text(p, sub):
    for r in p.runs:                       # single-run case (keeps formatting)
        if sub in r.text:
            r.text = r.text.replace(sub, '')
            return True
    full = ''.join(r.text for r in p.runs)  # spans runs -> rebuild into run[0]
    if sub in full and p.runs:
        for r in list(p.runs[1:]):
            r._element.getparent().remove(r._element)
        p.runs[0].text = full.replace(sub, '')
        return True
    return False


# --- 1) remove ALL separator lines (title + section headings) ---
borders = sum(drop_borders(p) for p in d.paragraphs)

# --- 2) remove bot-like English gloss parentheses ---
removed = []
for sub in [' (Twitter)', ' (Active Learning)', ' (uncertainty sampling)']:
    for p in d.paragraphs:
        if remove_text(p, sub):
            removed.append(sub.strip())
            break

d.save(str(DOC))
print("separator lines removed:", borders)
print("glosses removed:", removed)
print("\n--- title block now ---")
for p in d.paragraphs[:8]:
    if p.text.strip():
        print("  ", p.text.strip())
