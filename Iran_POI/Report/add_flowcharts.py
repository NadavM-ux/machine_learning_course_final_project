# -*- coding: utf-8 -*-
"""Surgical update: insert the Step-4 decision-flowchart subsection into the report.

Places a new 'שלב 4' subsection (3 flowchart images + accurate per-chart captions)
between the step-3 manual-labeling subsection and 'הפקת פיצ'רים', inside ב. שיטה.
Run ONCE (not idempotent).
"""
from pathlib import Path
from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

HERE = Path(__file__).resolve().parent
DOC = HERE / 'Iran_POI_Report.docx'
FLOW = HERE.parent / 'Step4_Decision_Flowcharts'
IMGS = [
    (FLOW / 'flowchart_1_target_population.png',
     "תרשים 1 — target_population (שיוך לאוכלוסיית היעד): (1) פרופיל מושהה/ריק/מוגן ללא מידע "
     "→ unknown. (2) אחרת, אם הביו, שם התצוגה או שם-המשתמש מזכירים Iran / Iranian / Persian / "
     "Persia או עיר איראנית → Target. (3) אחרת, אם שדה המיקום מצביע על עיר איראנית → Target. "
     "(4) אחרת, אם ציוצים אישיים מעידים שהאדם איראני → Target. (5) אחרת — אם קיימת זהות "
     "לא-איראנית ברורה → Not Target, ואם לא → Unknown."),
    (FLOW / 'flowchart_2_locals_vs_diaspora.png',
     "תרשים 2 — locals_vs_diaspora (מקומי מול תפוצות; רלוונטי רק כאשר target=1): (1) שדה המיקום "
     "מציג עיר איראנית → Local. (2) אחרת, שדה המיקום מציג עיר מחוץ לאיראן → diaspora. (3) אחרת, "
     "אם הביו/שם התצוגה מרמזים על תפוצה → diaspora. (4) אחרת, אם ציוצים אחרונים ממקמים בבירור "
     "את המשתמש בתוך איראן → local. (5) אחרת, אם ציוצים אחרונים ממקמים אותו בבירור מחוץ לאיראן "
     "→ diaspora, ואם לא → Unknown."),
    (FLOW / 'flowchart_3_person_vs_organization.png',
     "תרשים 3 — person_vs_organization (אדם מול ארגון): (1) פרופיל מושהה/ריק/פרטי → Unknown. "
     "(2) אחרת, אם שם התצוגה/שם-המשתמש מעידים בבירור על ארגון → organization. (3) אחרת, אם הביו "
     "מתאר מוסד → organization. (4) אחרת, אם תמונת הפרופיל היא לוגו/סמל-מותג/דמות לא-אנושית "
     "וגם הציוצים כתובים בקול מוסדי/רבים → organization. (5) אחרת, אם יש ראיה ברורה לאדם יחיד "
     "→ Person, ואם לא → Unknown."),
]

doc = Document(str(DOC))


def set_rtl(p):
    p._p.get_or_add_pPr().append(OxmlElement('w:bidi'))


anchor = next(p for p in doc.paragraphs if p.text.strip() == "הפקת פיצ'רים")

# 1) heading
h = anchor.insert_paragraph_before("תרשימי זרימה להחלטת התיוג (שלב 4)", style='Heading 2')
set_rtl(h)

# 2) intro
intro = anchor.insert_paragraph_before(
    "כדי לתעד באופן פורמלי את תהליך קבלת ההחלטות בתיוג הידני, בנינו שלושה תרשימי זרימה — "
    "אחד לכל משימת סיווג. כל תרשים מתאר את רצף השאלות (צומת-אחר-צומת) שלפיו הוכרע כל משתמש, "
    "כאשר כל עלה מוביל לאחד מהערכים 1/0/2. אותם תרשימים בדיוק שימשו גם כפרומפט למודל השפה "
    "בשלב 8-B, כך שהמסווג המאומן וה-LLM הוערכו על פי אותה לוגיקת החלטה. הקובץ המלא: "
    "user_labeling_decision_flows.pptx.")
set_rtl(intro)

# 3) three images, each followed by an accurate caption
for path, cap in IMGS:
    ip = anchor.insert_paragraph_before(style=None)
    ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ip.add_run().add_picture(str(path), width=Inches(6.3))

    cp = anchor.insert_paragraph_before(style=None)
    set_rtl(cp)
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cp.add_run(cap)
    r.italic = True
    r.font.size = Pt(9)

doc.save(str(DOC))
print("saved. paragraphs:", len(doc.paragraphs), "| tables:", len(doc.tables))
# count embedded images now
imgs = sum(1 for rel in doc.part.rels.values() if 'image' in rel.reltype)
print("embedded images:", imgs)
