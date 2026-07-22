# -*- coding: utf-8 -*-
"""Gap 2: step-7 two-band additional-labeling report. Additive only.
Adds a table (counts + manual-vs-model agreement per band) before 'שלב 7+'."""
from pathlib import Path
import pandas as pd
from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

HERE = Path(__file__).resolve().parent
SAMP = HERE.parent / 'Classification' / 'Step7_Stopping_Criteria' / \
    'stopping_criteria_probability_group_samples_for_manual_labeling.csv'
DOC = HERE / 'Iran_POI_Report.docx'

df = pd.read_csv(SAMP)
df['tp'] = pd.to_numeric(df['target_population'], errors='coerce').astype('Int64')
df['pc'] = pd.to_numeric(df['predicted_class'], errors='coerce').astype('Int64')

ORDER = ['uncertain [0.45-0.55]', 'confident [0.85-0.95]']
HEB = {'uncertain [0.45-0.55]': 'אי-ודאות [0.45-0.55]',
       'confident [0.85-0.95]': 'ביטחון גבוה [0.85-0.95]'}
rows = []
print("=== SECTION-2 RESULT (manual vs model agreement per band) ===")
for g in ORDER:
    sub = df[df['prob_group'] == g]
    n = len(sub)
    t1 = int((sub['tp'] == 1).sum()); t0 = int((sub['tp'] == 0).sum()); t2 = int((sub['tp'] == 2).sum())
    agree = float((sub['tp'] == sub['pc']).mean()) * 100
    rows.append([HEB[g], n, t1, t0, t2, f"{agree:.0f}%"])
    print(f"  {g:22} n={n} | target={t1} non={t0} unknown={t2} | agreement={agree:.0f}%")

doc = Document(str(DOC))


def rtl(p):
    p._p.get_or_add_pPr().append(OxmlElement('w:bidi')); return p


anchor = next(p for p in doc.paragraphs if p.text.strip().startswith('שלב 7+'))
capp = anchor.insert_paragraph_before(
    "דו\"ח תיוג נוסף - שתי קבוצות ההסתברות (שלב 7): 40 משתמשים שתויגו ידנית, 20 מכל קבוצה.",
    style=None)
rtl(capp)
cursor = capp._p

t = doc.add_table(rows=1, cols=6); t.style = 'Table Grid'
for c, h in zip(t.rows[0].cells,
                ['קבוצת הסתברות', 'n', 'target (1)', 'non-target (0)', 'unknown (2)', 'אחוז הסכמה למודל']):
    c.text = h
for r in rows:
    cells = t.add_row().cells
    for c, v in zip(cells, r):
        c.text = str(v)
cursor.addnext(t._tbl); cursor = t._tbl

note = doc.add_paragraph(); rtl(note)
run = note.add_run(
    "ההסכמה בין התיוג הידני לתחזית המודל עולה מ-55% בקבוצת אי-הוודאות ל-90% בקבוצת הביטחון "
    "הגבוה. כלומר רמת הביטחון של המודל מכוילת: כשהוא בטוח הוא צודק ברוב המכריע, וכשהוא לא בטוח "
    "אי-הוודאות אמיתית (40% מקרב הלא-ודאיים סומנו unknown גם בידי מתייג אנושי). ממצא זה מאשש את "
    "החלטת העצירה - תיוג נוסף באזור אי-הוודאות אינו צפוי לשפר את המודל.")
run.font.size = Pt(11)
cursor.addnext(note._p)

doc.save(str(DOC))
print("\nsaved. tables:", len(doc.tables), "| paragraphs:", len(doc.paragraphs))
