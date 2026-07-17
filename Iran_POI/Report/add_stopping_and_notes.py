# -*- coding: utf-8 -*-
"""Approved edits: rewrite the stopping paragraph on the winning model, add the
winning-model definition (section ב'), footnote T6, and a mean_AUC note under T7."""
from pathlib import Path
from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Pt

DOC = Path(__file__).resolve().parent / 'Iran_POI_Report.docx'
doc = Document(str(DOC))


def rtl(p):
    p._p.get_or_add_pPr().append(OxmlElement('w:bidi')); return p


def set_text(p, text):
    for r in list(p.runs[1:]):
        r._element.getparent().remove(r._element)
    (p.runs[0] if p.runs else p.add_run('')).text = text


def new_note(text, size=9, italic=True):
    p = doc.add_paragraph(); rtl(p)
    r = p.add_run(text); r.font.size = Pt(size); r.italic = italic
    return p


def find_table(headtext):
    for t in doc.tables:
        if headtext in " ".join(c.text for c in t.rows[0].cells):
            return t
    return None


# ---- 1) rewrite the stopping paragraph (winning model, +0.0009 / -0.0051) ----
sp = next(p for p in doc.paragraphs if 'החלטת העצירה' in p.text)
set_text(sp,
    "העצירה נשענה על שלושה סימנים: (א) ה-AUC של המודל המנצח על target_population ירד ב-0.0051 "
    "בין איטרציה 6 ל-7; (ב) ה-accuracy של אותו מודל נותר שטוח (+0.0009), הרבה מתחת לסף 0.5%, "
    "כלומר אינו שיפור מדיד; (ג) תיוג ידני של 20 משתמשים באזור אי-הוודאות הראה שב-40% מהם גם "
    "מתייג אנושי סימן unknown, כלומר אי-הוודאות אמיתית, ותיוג נוסף באזור זה אינו צפוי לשפר את המודל.")

# ---- 2) winning-model definition in section ב' (before the AL heading) ----
anchor = next(p for p in doc.paragraphs if p.text.strip().startswith('למידה אקטיבית וקריטריון'))
wm = anchor.insert_paragraph_before(
    "בחירת המודל המנצח בכל איטרציה נעשתה לפי ה-AUC הגבוה ביותר מבין ניסויי K-Fold+balanced בלבד. "
    "הרצות LOOCV בוצעו ותועדו כנדרש, אך הוחרגו מבחירת המנצח: בהיעדר הסתברויות רציפות ברמת ה-fold "
    "הן מייצרות AUC מנוון (עד ~1.0), ובחירה לפיהן הייתה מטעה.", style=None)
rtl(wm)

# ---- 3) footnote under T6 (enrichment table) ----
t6 = find_table('Acc בסיס')
fn6 = new_note("בטבלה זו עמודת ה-Acc היא ה-accuracy המרבי העצמאי, ולא בהכרח של מודל ה-max-AUC - "
               "הקונבנציה שונה מיתר הטבלאות, ונועדה להראות שההעשרה מרימה את שני המדדים בנפרד.")
t6._tbl.addnext(fn6._p)

# ---- plan-3: mean_AUC note under T7 (summary table) ----
t7 = find_table('Accuracy ממוצע')
fn7 = new_note("ה-mean_AUC נמוך יחסית כי הוא ממצע גם ריצות LOOCV (המייצרות AUC מנוון) וריצות 3-מחלקות "
               "והמשימה המנוונת locals_vs_diaspora; המדד הקובע להחלטת העצירה הוא ה-AUC של המודל המנצח "
               "על target_population.")
t7._tbl.addnext(fn7._p)

doc.save(str(DOC))
print("done. tables:", len(doc.tables), "| paragraphs:", len(doc.paragraphs))
# echo the edits
d2 = Document(str(DOC))
print("\nSTOPPING:", next(p.text for p in d2.paragraphs if p.text.strip().startswith('העצירה נשענה'))[:80], "...")
print("SECTION ב':", next(p.text for p in d2.paragraphs if 'בחירת המודל המנצח' in p.text)[:70], "...")
print("no arrows:", not any('→' in p.text for p in d2.paragraphs), "| no em-dash:", not any('—' in p.text for p in d2.paragraphs))
